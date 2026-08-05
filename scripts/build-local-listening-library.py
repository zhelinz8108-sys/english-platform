#!/usr/bin/env python3
"""Build the complete local listening catalogue used by demo and production import."""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import logging
import math
import re
import time
import urllib.parse
import urllib.request
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import UTC, datetime
from functools import lru_cache
from pathlib import Path
from typing import Any

from docx import Document
from mutagen import File as MutagenFile
from pypdf import PdfReader

try:
    import eng_to_ipa as english_ipa
except ImportError:  # IPA is helpful but catalogue generation can continue without it.
    english_ipa = None

logging.getLogger("pypdf").setLevel(logging.ERROR)


AUDIO_EXTENSIONS = {".mp3", ".mp4", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus", ".wma"}
LEGACY_AUDIO_EXTENSIONS = AUDIO_EXTENSIONS - {".mp4"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
TRANSCRIPT_EXTENSIONS = {".srt", ".lrc", ".txt", ".pdf", ".docx"}
DUPLICATE_SUFFIX = re.compile(r"\s*\(\d+\)$")
DATE_PATTERN = re.compile(r"(?<!\d)(20\d{6}|\d{6})(?!\d)")
BBC_YEAR_PATTERN = re.compile(r"(?<!\d)(20\d{2})(?!\d)")
MONTH_PATTERN = re.compile(r"(?<!\d)(20\d{4})(?!\d)")
SRT_TIMECODE_PATTERN = re.compile(
    r"^\s*\d{1,2}:\d{2}(?::\d{2})?[,.]\d{2,3}\s*-->\s*\d{1,2}:\d{2}(?::\d{2})?[,.]\d{2,3}.*$"
)
INLINE_TIMESTAMP_PATTERN = re.compile(
    r"^\s*(?:\[\d{1,2}:\d{2}(?::\d{2})?(?:[.:-]\d{1,3})?(?:\s*-\s*\d{1,2}:\d{2}(?::\d{2})?(?:[.:-]\d{1,3})?)?\]\s*)+"
)
CJK_PATTERN = re.compile(r"[\u3400-\u9fff]")
WORD_PATTERN = re.compile(r"\b[A-Za-z][A-Za-z'-]{3,}\b")
TERM_PATTERN = re.compile(
    r"^(?:[A-Za-z(][A-Za-z'()/-]*)(?: [A-Za-z(][A-Za-z'()/-]*){0,7}$"
)
VOCABULARY_MARKER = re.compile(r"(?im)^\s*VOCABULARY\s*$")
BBC_CHINESE_OVERRIDES = {
    "not authorised by people in authority": "未经当局授权",
    "not following the right route": "偏离正确路线",
    "involves a lot of people to do something": "需要大量人力来完成某事",
    "judging something by the way it looks": "从外观来评判某物",
    "is understandable and not a surprise": "合乎情理，并不令人意外",
    "it seems likely to be true (that); it makes sense (that) back to the dark ages": "按理说……很可能是真的；……是合乎情理的",
    "not follow what most people are doing; do the opposite of most people": "不随大流；与大多数人的做法相反",
    "noisy, uncontrolled outburst of anger": "吵闹且无法控制的愤怒爆发",
    "kept private and secret; information which is not shared with anyone": "保密的；不与任何人分享的信息",
    "not completely or entirely true": "不完全属实",
    "nonsense or not true": "胡说；不真实的说法",
    "is obviously true from the facts": "根据事实显而易见；合乎情理",
    "not a serious or significant problem": "不严重或不重要的问题",
    "not able to be played": "无法游玩的；玩不了的",
}
CONTEXT_CHINESE_OVERRIDES = {
    "Hana arigatou gozaimashita.": "Hana，非常感谢你。",
}
POS_ABBREVIATIONS = {
    "vt": "v.t.",
    "vi": "v.i.",
    "v": "v.",
    "n": "n.",
    "a": "adj.",
    "adj": "adj.",
    "na": "adj.",
    "s": "adj.",
    "ad": "adv.",
    "adv": "adv.",
    "prep": "prep.",
    "conj": "conj.",
    "pron": "pron.",
    "num": "num.",
    "art": "art.",
    "aux": "aux.",
    "int": "interj.",
    "interj": "interj.",
    "abbr": "abbr.",
}
POS_PATTERN = re.compile(
    r"^(vt|vi|adj|adv|prep|conj|pron|num|interj|abbr|aux|na|ad|art|int|v|n|a|s)(?=\.|\s|\[)\.?\s*(.*)$",
    flags=re.IGNORECASE,
)
VOCABULARY_STOP_WORDS = {
    "about",
    "after",
    "again",
    "also",
    "another",
    "because",
    "before",
    "being",
    "between",
    "could",
    "didn't",
    "doesn't",
    "during",
    "english",
    "every",
    "first",
    "going",
    "great",
    "hello",
    "little",
    "maybe",
    "minute",
    "people",
    "programme",
    "question",
    "really",
    "right",
    "should",
    "something",
    "that's",
    "their",
    "there",
    "these",
    "thing",
    "think",
    "those",
    "through",
    "today",
    "using",
    "very",
    "we're",
    "we've",
    "well",
    "what's",
    "where",
    "which",
    "while",
    "would",
    "you're",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--minute-earth-root", type=Path, required=True)
    parser.add_argument("--minute-earth-study-content", type=Path, required=True)
    parser.add_argument("--bbc-root", type=Path, required=True)
    parser.add_argument("--bbc-minute-root", type=Path)
    parser.add_argument("--voa-standard-root", type=Path)
    parser.add_argument("--scientific-american-root", type=Path)
    parser.add_argument("--short-wave-root", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument(
        "--dictionary",
        type=Path,
        help="ECDICT CSV. Defaults to Minute Earth/tmp/ecdict_repo/ecdict.csv.",
    )
    parser.add_argument(
        "--translation-cache",
        type=Path,
        help="Cached Chinese translations for BBC definitions and terms.",
    )
    parser.add_argument(
        "--context-translation-cache",
        type=Path,
        help="Cached Chinese translations for vocabulary contexts.",
    )
    parser.add_argument(
        "--transcript-overrides",
        type=Path,
        help="Verified speech-to-text overrides. Defaults beside the output library.",
    )
    parser.add_argument(
        "--translate-missing",
        action="store_true",
        help="Fill missing Chinese definitions and contexts through Google Translate and update the caches.",
    )
    parser.add_argument("--workers", type=int, default=8)
    return parser.parse_args()


def canonical_stem(path: Path) -> str:
    return DUPLICATE_SUFFIX.sub("", path.stem).strip()


def preferred_copy(paths: list[Path]) -> Path:
    return sorted(paths, key=lambda path: (bool(DUPLICATE_SUFFIX.search(path.stem)), len(path.name), path.name))[0]


def audio_duration(path: Path) -> int | None:
    try:
        audio = MutagenFile(path)
        if audio is None or audio.info is None:
            return None
        return round(float(audio.info.length))
    except Exception:
        return None


def stable_source_id(prefix: str, path: Path, root: Path, title: str) -> str:
    relative = path.relative_to(root).as_posix().casefold()
    digest = hashlib.sha1(relative.encode("utf-8")).hexdigest()[:12]
    return f"{prefix}-{digest}-{slug(title)[:36]}".rstrip("-")


def read_text_fallback(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-16", "gb18030", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def english_only_transcript(text: str) -> str:
    text = html.unescape(text).replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    kept: list[str] = []
    previous = ""
    for raw_line in text.splitlines():
        line = INLINE_TIMESTAMP_PATTERN.sub("", raw_line)
        line = re.sub(r"<[^>]+>", "", line)
        line = re.sub(r"\{\\[^}]+\}", "", line)
        line = re.sub(r"[ \t]+", " ", line).strip()
        if not line or line.isdigit() or SRT_TIMECODE_PATTERN.fullmatch(line):
            continue
        latin_words = re.findall(r"[A-Za-z][A-Za-z'’-]*", line)
        if CJK_PATTERN.search(line):
            english_segments = re.split(r"[\u3400-\u9fff]+", line)
            line = max(
                english_segments,
                key=lambda segment: len(re.findall(r"[A-Za-z][A-Za-z'’-]*", segment)),
            ).strip(" ：:，,。；;\t")
            latin_words = re.findall(r"[A-Za-z][A-Za-z'’-]*", line)
        if not latin_words or len(latin_words) < 3:
            continue
        normalized = line.casefold()
        if normalized in {
            "bbc learning english",
            "english in a minute",
            "this is not a word-for-word transcript",
        }:
            continue
        if normalized == previous:
            continue
        kept.append(line)
        previous = normalized
    return "\n".join(kept).strip()


@lru_cache(maxsize=None)
def extract_transcript(path: Path) -> tuple[str, int]:
    if path.suffix.casefold() in {".pdf", ".docx"}:
        text, _ = extract_document(path)
    else:
        text = read_text_fallback(path)
    transcript = english_only_transcript(text)
    return transcript, len(re.findall(r"\b[A-Za-z][A-Za-z'’-]*\b", transcript))


def normalized_file_key(value: str) -> str:
    value = DUPLICATE_SUFFIX.sub("", value.casefold())
    value = value.replace("&", " and ")
    return "".join(re.findall(r"[a-z0-9\u3400-\u9fff]+", value))


def same_parent_transcript(audio: Path, candidates: list[Path]) -> Path | None:
    siblings = [candidate for candidate in candidates if candidate.parent == audio.parent]
    if not siblings:
        return None
    direct = [candidate for candidate in siblings if candidate.stem.casefold() == audio.stem.casefold()]
    if direct:
        return preferred_transcript(direct)
    audio_key = normalized_file_key(audio.stem)
    exact = [candidate for candidate in siblings if normalized_file_key(candidate.stem) == audio_key]
    if len(exact) == 1:
        return preferred_transcript(exact)
    audio_date = normalized_date_token(audio.stem) or normalized_date_token(audio.parent.name)
    dated = [
        candidate
        for candidate in siblings
        if audio_date
        and (normalized_date_token(candidate.stem) or normalized_date_token(candidate.parent.name))
        == audio_date
    ]
    if len(dated) == 1:
        return dated[0]
    sibling_media = [path for path in audio.parent.iterdir() if path.is_file() and path.suffix.casefold() in AUDIO_EXTENSIONS]
    if len(siblings) == 1 and len(sibling_media) == 1:
        return siblings[0]
    return None


def preferred_transcript(paths: list[Path]) -> Path:
    priority = {".srt": 0, ".lrc": 1, ".txt": 2, ".pdf": 3, ".docx": 4}
    return sorted(paths, key=lambda path: (priority.get(path.suffix.casefold(), 9), len(path.name), path.name))[0]


def clean_pdf_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    cleaned: list[str] = []
    for line in lines:
        if re.fullmatch(r"Page \d+ of \d+", line, flags=re.IGNORECASE):
            continue
        if re.fullmatch(r"©?\s*bbclearningenglish\.com\s*\d{4}", line, flags=re.IGNORECASE):
            continue
        if line or (cleaned and cleaned[-1]):
            cleaned.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(cleaned)).strip()


@lru_cache(maxsize=None)
def extract_document(path: Path) -> tuple[str, int]:
    try:
        if path.suffix.casefold() == ".docx":
            document = Document(path)
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            text = "\n".join(blocks)
        else:
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        cleaned = clean_pdf_text(text)
        return cleaned, len(cleaned.split())
    except Exception:
        return "", 0


@lru_cache(maxsize=None)
def extract_document_layout(path: Path) -> str:
    try:
        if path.suffix.casefold() == ".docx":
            document = Document(path)
            return "\n\n".join(
                paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
            )
        reader = PdfReader(path)
        pages: list[str] = []
        for page in reader.pages[-2:]:
            try:
                pages.append(page.extract_text(extraction_mode="layout") or "")
            except Exception:
                pages.append(page.extract_text() or "")
        return "\n".join(pages)
    except Exception:
        return ""


@lru_cache(maxsize=None)
def american_ipa(term: str) -> str:
    if english_ipa is None:
        return ""
    try:
        value = english_ipa.convert(term).strip()
        if not value or "*" in value:
            return ""
        return f"/{value}/"
    except Exception:
        return ""


def normalize_text(value: str) -> str:
    replacements = {
        "\u00a0": " ",
        "\u200b": "",
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
    }
    normalized = html.unescape(value)
    for old, new in replacements.items():
        normalized = normalized.replace(old, new)
    return re.sub(r"\s+", " ", normalized).strip()


def parsed_chinese_definition(value: str) -> tuple[str, str]:
    definitions: list[tuple[str, str]] = []
    value = value.replace("\\r\\n", "\n").replace("\\n", "\n")
    for raw in value.splitlines():
        line = raw.strip()
        if not line:
            continue
        line = line.replace(",", "，").replace(";", "；")
        match = POS_PATTERN.match(line)
        if match:
            part_of_speech = POS_ABBREVIATIONS[match.group(1).casefold()]
            meaning = match.group(2).strip()
        else:
            part_of_speech = ""
            meaning = line
        if meaning:
            definitions.append((part_of_speech, meaning.rstrip("，； ")))
    parts = list(dict.fromkeys(part for part, _ in definitions if part))
    part_of_speech = " / ".join(parts)
    chinese = "；".join(f"{part} {meaning}".strip() for part, meaning in definitions)
    return part_of_speech, chinese


def dictionary_ipa(value: str) -> str:
    normalized = normalize_text(value).replace("ә", "ə").replace(":", "ː").strip("/ ")
    return f"/{normalized}/" if normalized else ""


def valid_vocabulary_term(value: str) -> bool:
    normalized = " ".join(value.split()).strip(" -")
    return (
        1 <= len(normalized) <= 80
        and bool(re.search(r"[A-Za-z]", normalized))
        and bool(TERM_PATTERN.fullmatch(normalized))
        and not re.search(r"^(?:bbc|minute english|page)$", normalized, flags=re.IGNORECASE)
    )


def extract_official_vocabulary(path: Path) -> list[dict[str, str]]:
    text = extract_document_layout(path)
    markers = list(VOCABULARY_MARKER.finditer(text))
    if not markers:
        return []
    segment = text[markers[-1].end() :]
    segment = re.split(
        r"(?im)^\s*(?:6 Minute English|bbclearningenglish\.com)", segment, maxsplit=1
    )[0]
    vocabulary: list[dict[str, str]] = []
    seen: set[str] = set()
    for block in re.split(r"\n\s*\n+", segment):
        lines = [" ".join(line.split()) for line in block.splitlines() if line.strip()]
        if len(lines) < 2:
            continue
        term = lines[0].strip(" -")
        definition = " ".join(lines[1:]).strip()
        key = term.casefold()
        if (
            key in seen
            or not valid_vocabulary_term(term)
            or len(definition.split()) < 2
            or len(definition) > 600
        ):
            continue
        seen.add(key)
        vocabulary.append(
            {
                "word": term,
                "ipa": american_ipa(term),
                "partOfSpeech": "",
                "definition": "",
                "englishDefinition": definition,
                "context": "",
            }
        )
    return vocabulary[:15]


def context_candidates(transcript: str) -> list[str]:
    candidates: list[str] = []
    # Some VOA transcripts join sentences without a space (``word.Next``).
    # Splitting before a following capital keeps decimals and abbreviations such
    # as ``3.5`` intact while still recovering usable context sentences.
    for raw in re.split(
        r"\n+|(?<=[.!?])\s+|(?<=[.!?])(?=[\"'“”‘’]?[A-Z])",
        transcript,
    ):
        sentence = " ".join(raw.split())
        if (
            8 <= len(sentence) <= 1200
            and "http" not in sentence.casefold()
            and "bbclearningenglish.com" not in sentence.casefold()
            and not sentence.casefold().startswith("vocabulary ")
        ):
            candidates.append(sentence)
    return candidates


def shortened_context(sentence: str, match: re.Match[str], limit: int = 280) -> str:
    if len(sentence) <= limit:
        return sentence
    left = max(0, match.start() - 105)
    right = min(len(sentence), match.end() + 145)
    if left:
        next_space = sentence.find(" ", left)
        left = next_space + 1 if next_space >= 0 else left
    if right < len(sentence):
        previous_space = sentence.rfind(" ", 0, right)
        right = previous_space if previous_space >= 0 else right
    excerpt = sentence[left:right].strip(" ,;:-")
    return f"{'…' if left else ''}{excerpt}{'…' if right < len(sentence) else ''}"


def vocabulary_search_terms(term: str) -> list[str]:
    normalized = normalize_text(term).casefold().strip(" -")
    variants = [normalized]
    variants.append(re.sub(r"\([^)]*\)", " ", normalized))
    variants.append(re.sub(r"[()]", "", normalized))
    expanded = list(variants)
    for value in variants:
        expanded.extend(part.strip() for part in re.split(r"\s*/\s*", value))
        expanded.append(re.sub(r"^(?:a|an|the|to|be)\s+", "", value))
    cleaned = [re.sub(r"\s+", " ", value).strip(" -") for value in expanded]
    return list(dict.fromkeys(value for value in cleaned if value))


def flexible_word_pattern(word: str) -> re.Pattern[str]:
    escaped = re.escape(word)
    if len(word) < 4 or not word.isalpha():
        return re.compile(rf"\b{escaped}\b", flags=re.IGNORECASE)
    if word.endswith("y") and len(word) > 4:
        forms = rf"{re.escape(word[:-1])}(?:y|ies|ied|ying)"
    elif word.endswith("e"):
        forms = rf"{escaped}(?:s|d)?|{re.escape(word[:-1])}ing"
    else:
        forms = rf"{escaped}(?:s|es|ed|ing|er|ers|ly)?"
    return re.compile(rf"\b(?:{forms})\b", flags=re.IGNORECASE)


def context_sentence(transcript: str, term: str) -> str:
    sentences = context_candidates(transcript)
    variants = vocabulary_search_terms(term)
    for variant in variants:
        pattern = re.compile(rf"\b{re.escape(variant)}\b", flags=re.IGNORECASE)
        for sentence in sentences:
            match = pattern.search(sentence)
            if match:
                return shortened_context(sentence, match)

    words = [
        word
        for word in re.findall(r"[a-z][a-z'-]+", " ".join(variants), flags=re.IGNORECASE)
        if len(word) >= 4 and word.casefold() not in VOCABULARY_STOP_WORDS
    ]
    for word in sorted(set(words), key=len, reverse=True):
        pattern = flexible_word_pattern(word.casefold())
        for sentence in sentences:
            match = pattern.search(sentence)
            if match:
                return shortened_context(sentence, match)

    irregular_forms = {
        "bacterium": "bacteria",
        "bind": "bound|binding",
        "bite": "bit|bitten|biting",
        "commit": "committed|committing",
        "compel": "compelled|compelling",
        "criterion": "criteria",
        "fungus": "fungi",
        "jog": "jogged|jogging",
        "larva": "larvae",
        "recur": "recurred|recurring",
        "spin": "spinning|spun",
    }
    irregular = irregular_forms.get(normalize_text(term).casefold())
    if irregular:
        pattern = re.compile(rf"\b(?:{irregular})\b", flags=re.IGNORECASE)
        for sentence in sentences:
            match = pattern.search(sentence)
            if match:
                return shortened_context(sentence, match)
    return ""


def add_context_vocabulary(items: list[dict[str, Any]]) -> None:
    transcripts = [str(item.get("transcript") or "") for item in items]
    document_frequency: Counter[str] = Counter()
    for transcript in transcripts:
        words = {
            match.group(0).casefold()
            for match in WORD_PATTERN.finditer(transcript)
            if match.group(0)[0].islower()
        }
        document_frequency.update(words)
    document_count = max(1, len([transcript for transcript in transcripts if transcript]))

    for item, transcript in zip(items, transcripts, strict=True):
        if item["vocabulary"] or not transcript:
            continue
        transcript_word_count = len(WORD_PATTERN.findall(transcript))
        if transcript_word_count <= 180:
            target_count = 8
        elif transcript_word_count <= 500:
            target_count = 10
        elif transcript_word_count <= 900:
            target_count = 14
        elif transcript_word_count <= 1_600:
            target_count = 16
        else:
            target_count = 20
        term_frequency: Counter[str] = Counter()
        for match in WORD_PATTERN.finditer(transcript):
            original = match.group(0)
            word = original.casefold().strip("'-")
            if (
                not original[0].islower()
                or word in VOCABULARY_STOP_WORDS
                or len(word) < 5
            ):
                continue
            term_frequency[word] += 1

        ranked: list[tuple[float, str]] = []
        for word, frequency in term_frequency.items():
            inverse_document_frequency = math.log(
                (document_count + 1) / (document_frequency[word] + 1)
            )
            score = inverse_document_frequency * (1 + math.log(frequency)) + min(len(word), 14) / 20
            ranked.append((score, word))
        ranked.sort(key=lambda entry: (-entry[0], entry[1]))

        vocabulary: list[dict[str, str]] = []
        for _, word in ranked[: max(80, target_count * 10)]:
            ipa = american_ipa(word)
            if english_ipa is not None and not ipa:
                continue
            sentence = context_sentence(transcript, word)
            if not sentence:
                continue
            vocabulary.append(
                {
                    "word": word,
                    "ipa": ipa,
                    "partOfSpeech": "",
                    "definition": "",
                    "englishDefinition": "",
                    "context": sentence,
                }
            )
            if len(vocabulary) == target_count:
                break
        item["vocabulary"] = vocabulary


def add_vocabulary_context(items: list[dict[str, Any]]) -> None:
    for item in items:
        transcript = str(item.get("transcript") or "")
        for entry in item["vocabulary"]:
            if not entry.get("context"):
                entry["context"] = context_sentence(transcript, str(entry["word"]))


def load_dictionary_rows(path: Path, terms: set[str]) -> dict[str, dict[str, str]]:
    if not path.exists():
        raise SystemExit(f"Missing ECDICT dictionary: {path}")
    rows: dict[str, dict[str, str]] = {}
    with path.open(encoding="utf-8", newline="") as file:
        for row in csv.DictReader(file):
            word = str(row.get("word") or "").casefold()
            if word in terms and row.get("translation"):
                rows[word] = row
    return rows


def load_translation_cache(path: Path) -> dict[str, dict[str, str]]:
    if not path.exists():
        return {"definitions": {}, "terms": {}}
    document = json.loads(path.read_text(encoding="utf-8"))
    return {
        "definitions": dict(document.get("definitions") or {}),
        "terms": dict(document.get("terms") or {}),
    }


def google_translate_batch(
    values: list[str], progress_label: str = "texts"
) -> dict[str, str]:
    def request_translation(payload: str) -> str:
        query = urllib.parse.urlencode(
            {"client": "gtx", "sl": "en", "tl": "zh-CN", "dt": "t", "q": payload}
        )
        url = f"https://translate.googleapis.com/translate_a/single?{query}"
        for attempt in range(3):
            try:
                request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(request, timeout=30) as response:
                    document = json.loads(response.read().decode("utf-8"))
                return "".join(segment[0] for segment in document[0] if segment[0])
            except Exception:
                if attempt == 2:
                    raise
                time.sleep(1.5 * (attempt + 1))
        return ""

    translations: dict[str, str] = {}
    batches: list[list[str]] = []
    current: list[str] = []
    current_length = 0
    for value in values:
        estimated = len(value) + 12
        if current and (len(current) >= 30 or current_length + estimated > 3800):
            batches.append(current)
            current = []
            current_length = 0
        current.append(value)
        current_length += estimated
    if current:
        batches.append(current)

    completed = 0
    for batch in batches:
        payload = "\n".join(f"[{index:02d}] {value}" for index, value in enumerate(batch))
        translated_text = request_translation(payload)
        matches = list(
            re.finditer(r"\[(\d{2})\]\s*(.*?)(?=\n?\[\d{2}\]|$)", translated_text, re.DOTALL)
        )
        translated_indexes: set[int] = set()
        for match in matches:
            index = int(match.group(1))
            if index >= len(batch):
                continue
            translated_indexes.add(index)
            source = batch[index]
            translations[source] = normalize_text(match.group(2)).rstrip("。")
        missing_indexes = [index for index in range(len(batch)) if index not in translated_indexes]
        if missing_indexes:
            print(
                f"Batch omitted {len(missing_indexes)} marker(s); translating those entries individually...",
                flush=True,
            )
            for index in missing_indexes:
                source = batch[index]
                translations[source] = normalize_text(request_translation(source)).rstrip("。")
        completed += len(batch)
        print(
            f"Translated {completed}/{len(values)} missing {progress_label}...",
            flush=True,
        )
        time.sleep(0.15)
    return translations


def load_context_translation_cache(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}
    document = json.loads(path.read_text(encoding="utf-8"))
    return dict(document.get("translations") or {})


def write_context_translation_cache(path: Path, translations: dict[str, str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(
            {
                "schemaVersion": 1,
                "generatedAt": datetime.now(UTC).isoformat(),
                "translations": translations,
            },
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        ),
        encoding="utf-8",
    )


def enrich_context_translations(
    items: list[dict[str, Any]],
    context_translation_cache_path: Path,
    translate_missing: bool,
) -> None:
    entries = [entry for item in items for entry in item["vocabulary"]]
    cache = load_context_translation_cache(context_translation_cache_path)
    cache.update(CONTEXT_CHINESE_OVERRIDES)
    contexts = sorted(
        {
            str(entry.get("context") or "")
            for entry in entries
            if entry.get("context")
        }
    )
    missing = [context for context in contexts if not cache.get(context)]

    if translate_missing and missing:
        checkpoint_size = 250
        for start in range(0, len(missing), checkpoint_size):
            checkpoint = missing[start : start + checkpoint_size]
            cache.update(
                google_translate_batch(checkpoint, progress_label="vocabulary contexts")
            )
            write_context_translation_cache(context_translation_cache_path, cache)
            completed = min(start + checkpoint_size, len(missing))
            print(
                f"Saved vocabulary context translations {completed}/{len(missing)}...",
                flush=True,
            )

    unresolved: list[str] = []
    for entry in entries:
        context = str(entry.get("context") or "")
        translation = cache.get(context, "") if context else ""
        entry["contextTranslation"] = translation
        if context and not translation:
            unresolved.append(context)

    if unresolved:
        print(
            f"Warning: {len(unresolved)} vocabulary rows lack Chinese context translations",
            flush=True,
        )


def infer_part_of_speech(term: str, english_definition: str) -> str:
    normalized = normalize_text(term).casefold()
    definition = normalize_text(english_definition).casefold()
    if " " in normalized or "(" in normalized or "/" in normalized:
        if re.search(r"^(?:\(?(?:to )?be\)?|to|go|get|give|have|make|take|turn)\b", normalized):
            return "phr.v."
        return "phr."
    if definition.startswith(("to ", "become ", "make ", "move ", "stop ", "do ")):
        return "v."
    if definition.startswith(("a ", "an ", "the ", "someone ", "something ", "person ", "people ")):
        return "n."
    if "used to describe" in definition or definition.startswith("describes "):
        return "adj."
    return "word"


def enrich_minute_earth_vocabulary(items: list[dict[str, Any]]) -> None:
    for item in items:
        for entry in item["vocabulary"]:
            part_of_speech, _ = parsed_chinese_definition(str(entry.get("definition") or ""))
            entry["partOfSpeech"] = part_of_speech or "word"
            entry["englishDefinition"] = ""
            entry["context"] = context_sentence(str(item.get("transcript") or ""), entry["word"])


def enrich_bbc_vocabulary(
    items: list[dict[str, Any]],
    dictionary_path: Path,
    translation_cache_path: Path,
    translate_missing: bool,
) -> None:
    entries = [entry for item in items for entry in item["vocabulary"]]
    terms = {str(entry["word"]).casefold() for entry in entries}
    dictionary = load_dictionary_rows(dictionary_path, terms)
    cache = load_translation_cache(translation_cache_path)
    cache["definitions"].update(BBC_CHINESE_OVERRIDES)
    missing_definitions = sorted(
        {
            str(entry.get("englishDefinition") or "")
            for entry in entries
            if entry.get("englishDefinition")
            and entry["englishDefinition"] not in cache["definitions"]
        }
    )
    missing_terms = sorted(
        {
            str(entry["word"])
            for entry in entries
            if not entry.get("englishDefinition")
            and entry["word"].casefold() not in dictionary
            and entry["word"] not in cache["terms"]
        }
    )
    if translate_missing:
        if missing_definitions:
            cache["definitions"].update(google_translate_batch(missing_definitions))
        if missing_terms:
            cache["terms"].update(google_translate_batch(missing_terms))
        translation_cache_path.parent.mkdir(parents=True, exist_ok=True)
        translation_cache_path.write_text(
            json.dumps(
                {
                    "schemaVersion": 1,
                    "generatedAt": datetime.now(UTC).isoformat(),
                    **cache,
                },
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            ),
            encoding="utf-8",
        )

    unresolved: list[str] = []
    for entry in entries:
        word = str(entry["word"])
        english_definition = str(entry.get("englishDefinition") or "")
        row = dictionary.get(word.casefold())
        dictionary_pos, dictionary_chinese = parsed_chinese_definition(
            str(row.get("translation") or "") if row else ""
        )
        if not entry.get("ipa") and row:
            entry["ipa"] = dictionary_ipa(str(row.get("phonetic") or ""))
        if english_definition:
            chinese = cache["definitions"].get(english_definition, "")
        else:
            chinese = dictionary_chinese or cache["terms"].get(word, "")
        if not chinese:
            unresolved.append(word)
            chinese = "中文释义待补充"
        entry["partOfSpeech"] = dictionary_pos or infer_part_of_speech(
            word, english_definition
        )
        entry["definition"] = chinese
        if not entry.get("context") and english_definition:
            entry["context"] = f"BBC Vocabulary: {word} — {english_definition}"

    if unresolved:
        print(f"Warning: {len(unresolved)} BBC vocabulary rows lack Chinese definitions")


def normalized_date_token(value: str) -> str | None:
    match = DATE_PATTERN.search(value)
    if not match:
        return None
    token = match.group(1)
    if len(token) == 8:
        return token
    year = int(token[:2])
    return ("20" if year < 50 else "19") + token


def date_token_for_path(path: Path) -> str | None:
    """Prefer the episode folder date over legacy dates embedded in media names."""
    return normalized_date_token(path.parent.name) or normalized_date_token(path.stem)


def bbc_year_for_path(path: Path, root: Path) -> int | None:
    """Keep the source library's year-folder structure even when filenames lack dates."""
    for part in path.relative_to(root).parts:
        match = BBC_YEAR_PATTERN.search(part)
        if match:
            return int(match.group(1))
    date_token = date_token_for_path(path)
    return int(date_token[:4]) if date_token else None


def vocabulary_key(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip().casefold())


def deduplicate_vocabulary(items: list[dict[str, Any]]) -> int:
    """Keep each vocabulary term only at its earliest occurrence in the catalogue."""
    seen: set[str] = set()
    removed = 0
    for item in items:
        unique_entries: list[dict[str, Any]] = []
        for entry in item["vocabulary"]:
            key = vocabulary_key(str(entry.get("word") or ""))
            if not key or key in seen:
                removed += 1
                continue
            seen.add(key)
            unique_entries.append(entry)
        item["vocabulary"] = unique_entries
    return removed


def humanize_title(value: str) -> str:
    value = DUPLICATE_SUFFIX.sub("", value)
    value = re.sub(r"【[^】]*】|\[[^\]]*\]", "", value)
    value = re.sub(r"^(?:6minute[_-]?)?\d{6,8}[_ -]*", "", value, flags=re.IGNORECASE)
    value = re.sub(r"(?:中英对照|对白|transcript|worksheet)$", "", value, flags=re.IGNORECASE)
    value = re.sub(
        r"(?:for[_ -]?web|download|audio|au[_ -]?bb|mp3)$",
        "",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", value)
    value = re.sub(r"[_-]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value or "BBC 6 Minute English"


def slug(value: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return normalized[:72] or "episode"


TITLE_STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "bbc",
    "english",
    "for",
    "is",
    "minute",
    "of",
    "the",
    "to",
}


def title_terms(value: str) -> set[str]:
    cleaned = humanize_title(value).casefold()
    return {
        word
        for word in re.findall(r"[a-z0-9]+", cleaned)
        if word not in TITLE_STOP_WORDS and not word.isdigit()
    }


def title_match_score(audio: Path, pdf: Path, *, inspect_content: bool) -> float:
    audio_source = audio.parent.name if normalized_date_token(audio.parent.name) else canonical_stem(audio)
    expected = title_terms(audio_source)
    if not expected:
        return 0
    filename_terms = title_terms(canonical_stem(pdf))
    score = len(expected & filename_terms) / len(expected)
    if score >= 0.6 or not inspect_content:
        return score
    transcript, _ = extract_document(pdf)
    first_page_terms = set(re.findall(r"[a-z0-9]+", transcript[:2500].casefold()))
    return max(score, len(expected & first_page_terms) / len(expected))


def build_minute_earth(root: Path, study_path: Path) -> list[dict[str, Any]]:
    document = json.loads(study_path.read_text(encoding="utf-8"))
    episodes = document.get("episodes", [])
    audio_candidates: dict[int, list[Path]] = {}
    for path in root.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in LEGACY_AUDIO_EXTENSIONS:
            continue
        match = re.match(r"^(\d{3})", path.name)
        if not match:
            continue
        sequence = int(match.group(1))
        if 1 <= sequence <= 270 and "人声分离测试" not in path.parts:
            audio_candidates.setdefault(sequence, []).append(path)

    items: list[dict[str, Any]] = []
    for episode in episodes:
        sequence = int(episode["sequence"])
        candidates = audio_candidates.get(sequence, [])
        audio_path = preferred_copy(candidates) if candidates else None
        transcript = str(episode.get("transcript", "")).strip()
        items.append(
            {
                "id": f"minute-earth-{sequence:03d}",
                "collection": "minute-earth",
                "sequence": sequence,
                "title": episode.get("title") or f"Minute Earth {sequence:03d}",
                "year": None,
                "publishedAt": None,
                "durationSeconds": episode.get("durationSeconds") or (audio_duration(audio_path) if audio_path else None),
                "sizeBytes": audio_path.stat().st_size if audio_path else 0,
                "audioPath": audio_path.relative_to(root).as_posix() if audio_path else None,
                "documentPath": None,
                "transcriptWordCount": episode.get("transcriptWordCount") or len(transcript.split()),
                "transcript": transcript,
                "vocabulary": episode.get("vocabulary", []),
            }
        )
    return items


def choose_bbc_document(audio: Path, documents: list[Path]) -> Path | None:
    audio_date = date_token_for_path(audio)
    audio_stem = canonical_stem(audio).casefold()
    candidate_classes: dict[Path, int] = {}
    split_scope = (
        audio.parent.parent
        if re.search(r"audio|mp3|音频", audio.parent.name, flags=re.IGNORECASE)
        else None
    )
    for document in documents:
        document_date = date_token_for_path(document)
        # Older BBC packages keep differently named audio and transcript files
        # together in one episode directory. Newer packages split audio and
        # documents into sibling directories, so their release date is the
        # reliable join key.
        if document.parent == audio.parent:
            candidate_classes[document] = 0
        elif audio_date and document_date == audio_date:
            candidate_classes[document] = 1
        elif split_scope is not None and document.parent.parent == split_scope:
            candidate_classes[document] = 2
    if not candidate_classes:
        return None

    filename_scores = {
        path: title_match_score(audio, path, inspect_content=False) for path in candidate_classes
    }
    best_filename_score = max(filename_scores.values(), default=0)
    match_scores = {
        path: (
            filename_scores[path]
            if candidate_classes[path] < 2 or best_filename_score >= 0.6
            else title_match_score(audio, path, inspect_content=True)
        )
        for path in candidate_classes
    }
    candidates = [
        path
        for path, candidate_class in candidate_classes.items()
        if candidate_class < 2 or match_scores[path] >= 0.6
    ]
    if not candidates:
        return None

    def rank(path: Path) -> tuple[int, float, int, int, bool, int, str]:
        name = canonical_stem(path).casefold()
        if name == audio_stem:
            priority = 0
        elif "对白" in path.stem or "transcrip" in name:
            priority = 1
        elif "中英对照" in path.stem or "中英文对照" in path.stem:
            priority = 2
        elif "worksheet" in name:
            priority = 4
        else:
            priority = 3
        return (
            candidate_classes[path],
            -match_scores[path],
            priority,
            0 if path.suffix.casefold() == ".pdf" else 1,
            bool(DUPLICATE_SUFFIX.search(path.stem)),
            len(path.name),
            path.name,
        )

    return sorted(candidates, key=rank)[0]


def build_bbc(root: Path, workers: int) -> list[dict[str, Any]]:
    audio_groups: dict[str, list[Path]] = {}
    documents = [
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.casefold() in DOCUMENT_EXTENSIONS
    ]
    for path in root.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in LEGACY_AUDIO_EXTENSIONS:
            continue
        relative_parent = path.parent.relative_to(root).as_posix().casefold()
        key = relative_parent + "/" + canonical_stem(path).casefold()
        audio_groups.setdefault(key, []).append(path)

    audios = [preferred_copy(group) for group in audio_groups.values()]
    audios.sort(
        key=lambda path: (
            bbc_year_for_path(path, root) or 9999,
            date_token_for_path(path) or "99999999",
            path.as_posix(),
        )
    )
    pairings = [(audio, choose_bbc_document(audio, documents)) for audio in audios]

    extracted: dict[Path, tuple[str, int]] = {}
    unique_documents = sorted({document for _, document in pairings if document is not None})
    with ThreadPoolExecutor(max_workers=max(1, workers)) as pool:
        futures = {
            pool.submit(extract_document, document): document for document in unique_documents
        }
        for future in as_completed(futures):
            extracted[futures[future]] = future.result()

    items: list[dict[str, Any]] = []
    used_ids: set[str] = set()
    for sequence, (audio, document) in enumerate(pairings, start=1):
        date_token = date_token_for_path(audio)
        source_year = bbc_year_for_path(audio, root)
        raw_title = audio.parent.name if normalized_date_token(audio.parent.name) else canonical_stem(audio)
        title = humanize_title(raw_title)
        base_id = f"bbc-{date_token or sequence:0>8}-{slug(title)}"
        item_id = base_id
        duplicate_index = 2
        while item_id in used_ids:
            item_id = f"{base_id}-{duplicate_index}"
            duplicate_index += 1
        used_ids.add(item_id)
        transcript, transcript_words = (
            extracted.get(document, ("", 0)) if document else ("", 0)
        )
        items.append(
            {
                "id": item_id,
                "collection": "bbc-6-minute-english",
                "sequence": sequence,
                "title": title,
                "year": source_year,
                "publishedAt": date_token,
                "durationSeconds": audio_duration(audio),
                "sizeBytes": audio.stat().st_size,
                "audioPath": audio.relative_to(root).as_posix(),
                "documentPath": document.relative_to(root).as_posix() if document else None,
                "transcriptWordCount": transcript_words,
                "transcript": transcript,
                "vocabulary": extract_official_vocabulary(document) if document else [],
            }
        )
    add_context_vocabulary(items)
    return items


def year_from_path(path: Path, root: Path) -> int | None:
    for part in path.relative_to(root).parts:
        match = BBC_YEAR_PATTERN.search(part)
        if match:
            return int(match.group(1))
    token = normalized_date_token(path.stem)
    return int(token[:4]) if token else None


def bbc_minute_episode_number(value: str) -> int | None:
    value = value.casefold()
    value = re.sub(r"^\d{6,8}", "", value)
    if re.fullmatch(r"\d{2,3}", value):
        return int(value)
    match = re.match(r"(?:eiam[_ -]*)?(\d{2,3})(?:bbc|[_ -])", value)
    return int(match.group(1)) if match else None


def bbc_minute_title(value: str) -> str:
    value = DUPLICATE_SUFFIX.sub("", value).strip(" \\_-")
    value = re.sub(r"^\d+-\d+\s*", "", value)
    value = re.sub(r"^\d{6,8}[_ -]*", "", value)
    value = re.sub(r"^eiam[_ -]*(?:\d{2,3}[_ -]*)?", "", value, flags=re.IGNORECASE)
    value = re.sub(r"^\d{2,3}bbc[_ -]*learning[_ -]*english[_ -]*", "", value, flags=re.IGNORECASE)
    value = re.sub(r"^\d{2,3}bbc[_ -]*english[_ -]*in[_ -]*a[_ -]*minute[_ -]*", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\.(?:en(?:-gb)?|eng)$", "", value, flags=re.IGNORECASE)
    value = re.sub(r"english\s+in\s+(?:a|m)\s+minute", "", value, flags=re.IGNORECASE)
    value = re.sub(r"[_-]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip(" ._-")
    return value or "BBC English in a Minute"


def bbc_minute_key(path: Path) -> str:
    if any("第三季" in part for part in path.parts):
        session = re.search(r"session[_ -]*(\d+)", path.stem, flags=re.IGNORECASE)
        if session:
            return f"season-3:{int(session.group(1)):02d}"
        if re.fullmatch(r"\d{1,2}", path.stem):
            return f"season-3:{int(path.stem):02d}"
    number = bbc_minute_episode_number(path.stem)
    if number is not None:
        return f"number:{number:03d}"
    title_key = normalized_file_key(bbc_minute_title(path.stem))
    title_aliases = {
        "5waystousehard": "5waysusehard",
        "5verbsfollowedbyverbing": "5verbsing",
        "goondoingvsgoontodowhatsthedifference": "goondoingtodo",
        "toovsenoughwhatsthedifference": "tooenough",
        "usingdifferenttensestobepolite": "politetenses",
        "verbstouseinsteadofwalk": "4wordsforwalk",
    }
    return f"title:{title_aliases.get(title_key, title_key)}"


def preferred_bbc_minute_document(paths: list[Path]) -> Path:
    def rank(path: Path) -> tuple[int, int, int, str]:
        suffix = path.suffix.casefold()
        official = bool(re.match(r"(?:\d{6,8}[_ -]*)?eiam", path.stem, flags=re.IGNORECASE))
        translated = bool(CJK_PATTERN.search(path.stem))
        return (
            0 if suffix == ".srt" else 1 if official else 2 if suffix == ".pdf" else 3,
            1 if translated else 0,
            len(path.name),
            path.name,
        )

    return sorted(paths, key=rank)[0]


def build_bbc_minute(root: Path, workers: int) -> list[dict[str, Any]]:
    media_groups: dict[str, list[Path]] = {}
    document_groups: dict[str, list[Path]] = {}
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        suffix = path.suffix.casefold()
        if suffix in AUDIO_EXTENSIONS:
            media_groups.setdefault(bbc_minute_key(path), []).append(path)
        elif suffix in TRANSCRIPT_EXTENSIONS:
            document_groups.setdefault(bbc_minute_key(path), []).append(path)

    def choose_media(paths: list[Path]) -> Path:
        priority = {".mp3": 0, ".m4a": 1, ".mp4": 2}
        return sorted(paths, key=lambda path: (priority.get(path.suffix.casefold(), 9), len(path.name), path.as_posix()))[0]

    pairings: list[tuple[str, Path, Path, Path]] = []
    for key, paths in media_groups.items():
        documents = document_groups.get(key, [])
        if not documents:
            continue
        audio = choose_media(paths)
        transcript_document = preferred_bbc_minute_document(documents)
        metadata_document = sorted(
            documents,
            key=lambda path: (
                normalized_date_token(path.stem) is None,
                year_from_path(path, root) is None,
                len(path.name),
            ),
        )[0]
        pairings.append((key, audio, transcript_document, metadata_document))

    pairings.sort(
        key=lambda pairing: (
            year_from_path(pairing[3], root) or 9999,
            bbc_minute_episode_number(pairing[1].stem) or 9999,
            bbc_minute_title(pairing[2].stem).casefold(),
        )
    )
    extracted: dict[Path, tuple[str, int]] = {}
    with ThreadPoolExecutor(max_workers=max(1, workers)) as pool:
        futures = {
            pool.submit(extract_transcript, document): document
            for _, _, document, _ in pairings
        }
        for future in as_completed(futures):
            extracted[futures[future]] = future.result()

    items: list[dict[str, Any]] = []
    for audio, transcript_document, metadata_document in (
        (audio, transcript_document, metadata_document)
        for _, audio, transcript_document, metadata_document in pairings
    ):
        transcript, transcript_words = extracted.get(transcript_document, ("", 0))
        if transcript_words < 8:
            continue
        sequence = len(items) + 1
        title = bbc_minute_title(transcript_document.stem)
        date_token = normalized_date_token(metadata_document.stem)
        source_year = year_from_path(metadata_document, root)
        items.append(
            {
                "id": stable_source_id("bbc-minute", audio, root, title),
                "collection": "bbc-english-in-a-minute",
                "sequence": sequence,
                "title": title,
                "year": source_year,
                "publishedAt": date_token,
                "durationSeconds": audio_duration(audio),
                "sizeBytes": audio.stat().st_size,
                "audioPath": audio.relative_to(root).as_posix(),
                "documentPath": transcript_document.relative_to(root).as_posix(),
                "transcriptWordCount": transcript_words,
                "transcript": transcript,
                "vocabulary": [],
            }
        )
    return items


def build_simple_transcript_collection(
    root: Path,
    *,
    collection: str,
    id_prefix: str,
    title_for_audio: Any,
    workers: int,
) -> list[dict[str, Any]]:
    audios = sorted(
        (path for path in root.rglob("*") if path.is_file() and path.suffix.casefold() in AUDIO_EXTENSIONS),
        key=lambda path: path.relative_to(root).as_posix().casefold(),
    )
    transcript_candidates = [
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.casefold() in TRANSCRIPT_EXTENSIONS
    ]
    pairings = [
        (audio, same_parent_transcript(audio, transcript_candidates))
        for audio in audios
    ]
    selected_documents = sorted({document for _, document in pairings if document is not None})
    extracted: dict[Path, tuple[str, int]] = {}
    with ThreadPoolExecutor(max_workers=max(1, workers)) as pool:
        futures = {
            pool.submit(extract_transcript, document): document for document in selected_documents
        }
        for future in as_completed(futures):
            extracted[futures[future]] = future.result()

    prepared: list[tuple[Path, Path, str, int]] = []
    for audio, document in pairings:
        if document is None:
            continue
        transcript, words = extracted.get(document, ("", 0))
        if words < 8:
            continue
        prepared.append((audio, document, transcript, words))

    prepared.sort(
        key=lambda pairing: (
            normalized_date_token(pairing[0].stem)
            or normalized_date_token(pairing[0].parent.name)
            or "99999999",
            pairing[0].relative_to(root).as_posix().casefold(),
        )
    )
    items: list[dict[str, Any]] = []
    for sequence, (audio, document, transcript, transcript_words) in enumerate(prepared, start=1):
        title = str(title_for_audio(audio))
        date_token = normalized_date_token(audio.stem) or normalized_date_token(audio.parent.name)
        source_year = year_from_path(audio, root)
        items.append(
            {
                "id": stable_source_id(id_prefix, audio, root, title),
                "collection": collection,
                "sequence": sequence,
                "title": title,
                "year": source_year,
                "publishedAt": date_token,
                "durationSeconds": audio_duration(audio),
                "sizeBytes": audio.stat().st_size,
                "audioPath": audio.relative_to(root).as_posix(),
                "documentPath": document.relative_to(root).as_posix(),
                "transcriptWordCount": transcript_words,
                "transcript": transcript,
                "vocabulary": [],
            }
        )
    return items


def short_wave_title(path: Path) -> str:
    value = re.sub(r"^\d{6}[_ -]*", "", DUPLICATE_SUFFIX.sub("", path.stem))
    return re.sub(r"\s+", " ", value.replace("_", " ")).strip() or "NPR Short Wave"


def scientific_american_title(path: Path) -> str:
    date_token = normalized_date_token(path.stem)
    category_match = re.search(r"sa[_ -]*(science|tech|earth|mind|health)", path.stem, flags=re.IGNORECASE)
    if date_token:
        date_label = f"{date_token[:4]}-{date_token[4:6]}-{date_token[6:8]}"
        category = category_match.group(1).title() if category_match else "Science"
        return f"Scientific American 60-Second · {date_label} · {category}"
    return re.sub(r"\s+", " ", DUPLICATE_SUFFIX.sub("", path.stem).replace("_", " ")).strip()


def apply_transcript_overrides(items: list[dict[str, Any]], path: Path) -> int:
    if not path.is_file():
        return 0
    document = json.loads(path.read_text(encoding="utf-8"))
    overrides = document.get("items")
    if not isinstance(overrides, list):
        raise SystemExit(f"Transcript override file has no items array: {path}")
    item_by_id = {str(item["id"]): item for item in items}
    applied = 0
    for override in overrides:
        if not isinstance(override, dict):
            raise SystemExit(f"Transcript override entry must be an object: {path}")
        source_id = str(override.get("id") or "").strip()
        transcript = str(override.get("transcript") or "").strip()
        item = item_by_id.get(source_id)
        if item is None:
            raise SystemExit(f"Transcript override references unknown source: {source_id}")
        if len(transcript.split()) < 20:
            raise SystemExit(f"Transcript override is too short for {source_id}")
        item["transcript"] = transcript
        item["transcriptWordCount"] = len(transcript.split())
        applied += 1
    return applied


def main() -> None:
    args = parse_args()
    minute_items = build_minute_earth(args.minute_earth_root.resolve(), args.minute_earth_study_content.resolve())
    bbc_items = build_bbc(args.bbc_root.resolve(), args.workers)
    bbc_minute_items = (
        build_bbc_minute(args.bbc_minute_root.resolve(), args.workers)
        if args.bbc_minute_root
        else []
    )
    voa_items = (
        build_simple_transcript_collection(
            args.voa_standard_root.resolve(),
            collection="voa-standard-english",
            id_prefix="voa-standard",
            title_for_audio=lambda path: path.stem,
            workers=args.workers,
        )
        if args.voa_standard_root
        else []
    )
    scientific_american_items = (
        build_simple_transcript_collection(
            args.scientific_american_root.resolve(),
            collection="scientific-american-60-second",
            id_prefix="scientific-american",
            title_for_audio=scientific_american_title,
            workers=args.workers,
        )
        if args.scientific_american_root
        else []
    )
    short_wave_items = (
        build_simple_transcript_collection(
            args.short_wave_root.resolve(),
            collection="short-wave",
            id_prefix="short-wave",
            title_for_audio=short_wave_title,
            workers=args.workers,
        )
        if args.short_wave_root
        else []
    )
    transcript_overrides_path = (
        args.transcript_overrides.resolve()
        if args.transcript_overrides
        else args.output.resolve().with_name("listening-transcript-overrides.json")
    )
    applied_transcript_overrides = apply_transcript_overrides(
        minute_items + bbc_items,
        transcript_overrides_path,
    )
    dictionary_path = (
        args.dictionary.resolve()
        if args.dictionary
        else (args.minute_earth_root.resolve() / "tmp" / "ecdict_repo" / "ecdict.csv")
    )
    translation_cache_path = (
        args.translation_cache.resolve()
        if args.translation_cache
        else args.output.resolve().with_name("bbc-vocabulary-translations.json")
    )
    context_translation_cache_path = (
        args.context_translation_cache.resolve()
        if args.context_translation_cache
        else args.output.resolve().with_name("listening-context-translations.json")
    )
    add_vocabulary_context(minute_items)
    add_vocabulary_context(bbc_items)
    enrich_minute_earth_vocabulary(minute_items)
    enrich_bbc_vocabulary(
        bbc_items,
        dictionary_path,
        translation_cache_path,
        args.translate_missing,
    )
    removed_bbc_vocabulary = deduplicate_vocabulary(bbc_items)
    enrich_context_translations(
        minute_items + bbc_items,
        context_translation_cache_path,
        args.translate_missing,
    )
    collection_definitions = [
        {
            "id": "bbc-english-in-a-minute",
            "label": "BBC 一分钟英语",
            "description": "一分钟语法与用法短讲，适合建立基础听辨与高频表达。",
            "difficulty": "A2",
            "audience": "初一–初二",
            "rank": 1,
            "count": len(bbc_minute_items),
        },
        {
            "id": "bbc-6-minute-english",
            "label": "BBC 6 Minute English",
            "description": "BBC 六分钟英语，含音频、原版对话稿和重点词汇。",
            "difficulty": "B1–B2",
            "audience": "初三优秀生–高中",
            "rank": 2,
            "count": len(bbc_items),
        },
        {
            "id": "voa-standard-english",
            "label": "VOA 常速英语新闻",
            "description": "常速国际新闻报道，配套英文逐字稿，训练真实新闻语速。",
            "difficulty": "B2",
            "audience": "高一–高二",
            "rank": 3,
            "count": len(voa_items),
        },
        {
            "id": "minute-earth",
            "label": "MinuteEarth",
            "description": "科学与地球主题短篇，含音频、英文原文和 TOEFL/SAT 词汇。",
            "difficulty": "B2",
            "audience": "高中；配合画面更易理解",
            "rank": 4,
            "count": len(minute_items),
        },
        {
            "id": "scientific-american-60-second",
            "label": "科学美国人 60 秒",
            "description": "一分钟科学新闻与研究解读，语速快、信息密度高。",
            "difficulty": "B2+–C1",
            "audience": "高中优秀生–大学",
            "rank": 5,
            "count": len(scientific_american_items),
        },
        {
            "id": "short-wave",
            "label": "Short Wave",
            "description": "NPR 科学播客，包含自然对话、采访与完整英文逐字稿。",
            "difficulty": "B2+–C1",
            "audience": "高中优秀生–大学",
            "rank": 6,
            "count": len(short_wave_items),
        },
    ]
    all_items = (
        bbc_minute_items
        + bbc_items
        + voa_items
        + minute_items
        + scientific_american_items
        + short_wave_items
    )
    output = {
        "schemaVersion": 6,
        "generatedAt": datetime.now(UTC).isoformat(),
        "collections": collection_definitions,
        "items": all_items,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(output, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")
    paired_bbc = sum(1 for item in bbc_items if item["documentPath"] and item["transcript"])
    vocabulary_bbc = sum(1 for item in bbc_items if item["vocabulary"])
    context_minute = sum(
        1 for item in minute_items for entry in item["vocabulary"] if entry["context"]
    )
    context_bbc = sum(
        1 for item in bbc_items for entry in item["vocabulary"] if entry["context"]
    )
    print(f"Minute Earth: {len(minute_items)} items")
    print(f"BBC: {len(bbc_items)} unique audio items, {paired_bbc} transcripts extracted")
    print(f"BBC English in a Minute: {len(bbc_minute_items)} paired items")
    print(f"VOA Standard English: {len(voa_items)} paired items")
    print(f"Scientific American 60-Second: {len(scientific_american_items)} paired items")
    print(f"Short Wave: {len(short_wave_items)} paired items")
    print(f"BBC vocabulary: {vocabulary_bbc} items")
    print(f"BBC vocabulary deduplicated: {removed_bbc_vocabulary} repeated cards removed")
    print(f"Transcript overrides applied: {applied_transcript_overrides}")
    print(f"Vocabulary contexts: Minute Earth {context_minute}, BBC {context_bbc}")
    print(f"Output: {args.output} ({args.output.stat().st_size / 1024 / 1024:.1f} MB)")


if __name__ == "__main__":
    main()

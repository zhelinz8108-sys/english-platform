#!/usr/bin/env python3
"""Export reviewed website listening content as a navigable DOCX workbook."""

from __future__ import annotations

import argparse
import json
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TEAL = "246B64"
TEAL_LIGHT = "E2F0EC"
GOLD = "C8821D"
GOLD_LIGHT = "FFF4DC"
INK = "222A2D"
MUTED = "687579"
RULE = "D7E1DE"
WHITE = "FFFFFF"
READY_STATUSES = {"reviewed", "adjudicated", "approved"}

TYPE_LABELS = {
    "main_idea": "主旨",
    "detail": "细节",
    "rhetorical_purpose": "修辞作用",
    "inference": "推断",
    "organization": "组织方式",
    "next_content": "后续内容",
    "prediction": "后续内容",
}
DIFFICULTY_LABELS = {"low": "基础", "medium": "中等", "high": "较难"}
COLLECTION_LABELS = {
    "minute-earth": "Minute Earth",
    "bbc-6-minute-english": "BBC 6 Minute English",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--library",
        type=Path,
        default=Path("apps/web/data/listening-library.json"),
    )
    parser.add_argument(
        "--question-bank",
        type=Path,
        default=Path("apps/web/data/toefl-academic-listening-questions/question-bank.json"),
    )
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument(
        "--limit",
        type=int,
        help="Export only the first N ready sets (for layout smoke tests).",
    )
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def mark_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_paragraph_shading(paragraph: Any, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph: Any, color: str, size: int = 16, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(
    style: Any,
    *,
    size: float,
    color: str = INK,
    bold: bool = False,
    before: float = 0,
    after: float = 6,
    line: float = 1.25,
) -> None:
    set_run_font(style.element.get_or_add_rPr(), size=size) if False else None
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def add_style(document: Document, name: str, *, base: str = "Normal", **kwargs: Any) -> Any:
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(name, 1)
    style.base_style = styles[base]
    configure_style(style, **kwargs)
    return style


def add_numbering_definition(document: Document, fmt: str, level_text: str, left: int, hanging: int) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def new_numbering_instance(document: Document, abstract_id: int) -> int:
    numbering = document.part.numbering_part.element
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph: Any, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)


def add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_section_label(document: Document, text: str, *, answer: bool = False) -> Any:
    paragraph = document.add_paragraph(style="Section Label")
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=GOLD if answer else TEAL, bold=True)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def evidence_region(start: int, transcript_length: int) -> tuple[str, int]:
    if transcript_length <= 0:
        return "位置未知", 0
    ratio = start / transcript_length
    label = "开头" if ratio < 1 / 3 else "中段" if ratio < 2 / 3 else "结尾"
    return label, round(ratio * 100)


def setup_document(document: Document) -> tuple[int, int]:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(document.styles["Normal"], size=11, color=INK, after=6, line=1.25)
    configure_style(document.styles["Subtitle"], size=14, color=MUTED, after=18, line=1.15)
    configure_style(document.styles["Heading 1"], size=18, color=TEAL, bold=True, before=18, after=10)
    configure_style(document.styles["Heading 2"], size=15, color=TEAL, bold=True, before=14, after=7)
    configure_style(document.styles["Heading 3"], size=13, color=INK, bold=True, before=10, after=5)
    for heading in ("Heading 1", "Heading 2", "Heading 3"):
        document.styles[heading].paragraph_format.keep_with_next = True

    add_style(document, "Cover Title", size=28, color=TEAL, bold=True, after=8, line=1.0)
    add_style(document, "Kicker", size=9, color=GOLD, bold=True, after=7, line=1.0)
    add_style(document, "Metadata", size=9, color=MUTED, after=4, line=1.15)
    add_style(document, "Section Label", size=11, color=TEAL, bold=True, before=10, after=5, line=1.0)
    add_style(document, "Transcript", size=10.5, color=INK, after=8, line=1.25)
    add_style(document, "Question Meta", size=8.5, color=TEAL, bold=True, before=7, after=2, line=1.0)
    add_style(document, "Question Prompt", size=11, color=INK, bold=True, after=5, line=1.2)
    add_style(document, "Option", size=10.5, color=INK, after=3, line=1.15)
    add_style(document, "Answer Lead", size=11, color=INK, bold=True, before=7, after=4, line=1.15)
    add_style(document, "Answer Body", size=10.5, color=INK, after=5, line=1.2)
    add_style(document, "Evidence", size=9.5, color=INK, after=5, line=1.15)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("AURELIS  ·  TOEFL ACADEMIC LISTENING PRACTICE")
    set_run_font(run, size=8.5, color=TEAL, bold=True)
    footer = section.footer
    add_page_field(footer.paragraphs[0])

    decimal_abstract = add_numbering_definition(document, "decimal", "%1.", 430, 300)
    option_abstract = add_numbering_definition(document, "upperLetter", "%1.", 500, 280)
    return decimal_abstract, option_abstract


def add_cover(
    document: Document,
    sets: list[dict[str, Any]],
    items_by_id: dict[str, dict[str, Any]],
) -> None:
    document.add_paragraph("AURELIS · LISTENING WORKBOOK", style="Kicker")
    document.add_paragraph("听力原文、题目与答案解析", style="Cover Title")
    document.add_paragraph(
        "Website-reviewed snapshot · TOEFL-style Academic Listening Practice",
        style="Subtitle",
    )
    counts = Counter(question_set["collection"] for question_set in sets)
    total_words = sum(items_by_id[item["sourceId"]].get("transcriptWordCount", 0) for item in sets)
    table = document.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    mark_table_header(table.rows[0])
    labels = ["已审定听力", "听力题目", "英文原文词数"]
    values = [f"{len(sets):,} 集", f"{len(sets) * 4:,} 题", f"{total_words:,} 词"]
    for index, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, TEAL_LIGHT)
        set_cell_margins(cell, top=120, bottom=55)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(labels[index])
        set_run_font(run, size=9, color=TEAL, bold=True)
    for index, cell in enumerate(table.rows[1].cells):
        set_cell_shading(cell, "F7FAF9")
        set_cell_margins(cell, top=70, bottom=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(values[index])
        set_run_font(run, size=18, color=INK, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("收录范围")
    set_run_font(run, size=12, color=TEAL, bold=True)
    range_text = (
        f"Minute Earth {counts.get('minute-earth', 0)} 集；"
        f"BBC 6 Minute English {counts.get('bbc-6-minute-english', 0)} 集。"
    )
    p = document.add_paragraph(range_text)
    p.style = document.styles["Normal"]
    p = document.add_paragraph(
        "每集先呈现完整英文原文与四道英文四选一题；四题之后集中呈现答案、中文解析、四个选项理由，以及证据原句在原文中的位置。",
    )
    p.style = document.styles["Normal"]
    p = document.add_paragraph(
        f"生成时间：{datetime.now().astimezone().strftime('%Y-%m-%d %H:%M')}；内容取自网站当前已审定题库快照。",
        style="Metadata",
    )
    set_paragraph_shading(p, GOLD_LIGHT)
    set_paragraph_left_border(p, GOLD, size=14, space=7)
    document.add_page_break()


def add_overview(document: Document, sets: list[dict[str, Any]], items_by_id: dict[str, dict[str, Any]]) -> None:
    document.add_paragraph("使用说明与内容导航", style="Heading 1")
    document.add_paragraph(
        "Word 中可打开“导航窗格”，按资料库、年份和集数标题快速跳转。题目答案统一放在每集四道题之后，适合先遮住答案完成练习，再核对解析。"
    )
    by_collection_year: dict[tuple[str, str], int] = defaultdict(int)
    for question_set in sets:
        item = items_by_id[question_set["sourceId"]]
        year = str(item.get("year") or "系列")
        by_collection_year[(question_set["collection"], year)] += 1
    for collection in ("minute-earth", "bbc-6-minute-english"):
        total = sum(count for (key, _), count in by_collection_year.items() if key == collection)
        if not total:
            continue
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{COLLECTION_LABELS[collection]} · {total} 集")
        set_run_font(run, size=11, color=TEAL, bold=True)
        years = sorted(
            ((year, count) for (key, year), count in by_collection_year.items() if key == collection),
            key=lambda pair: (pair[0] == "系列", pair[0]),
        )
        detail = "、".join(f"{year}：{count} 集" for year, count in years)
        document.add_paragraph(detail, style="Metadata")
    document.add_page_break()


def add_episode(
    document: Document,
    item: dict[str, Any],
    question_set: dict[str, Any],
    decimal_abstract: int,
    option_abstract: int,
    *,
    page_break_before: bool,
) -> None:
    title = f"{int(item.get('sequence') or 0):03d}. {item['title']}"
    heading = document.add_paragraph(title, style="Heading 3")
    heading.paragraph_format.page_break_before = page_break_before
    profile = "Academic Talk" if question_set["collection"] == "minute-earth" else "Academic Discussion"
    duration = item.get("durationSeconds") or 0
    minutes, seconds = divmod(int(duration), 60)
    metadata = (
        f"{COLLECTION_LABELS[question_set['collection']]} · {profile} · "
        f"{minutes}:{seconds:02d} · {item.get('transcriptWordCount', 0)} 词 · "
        f"4 题 · {question_set['status']}"
    )
    document.add_paragraph(metadata, style="Metadata")

    add_section_label(document, "英文原文  /  TRANSCRIPT")
    transcript = (item.get("transcript") or "").strip()
    paragraphs = [part.strip() for part in transcript.split("\n") if part.strip()] or [transcript]
    for text in paragraphs:
        document.add_paragraph(text, style="Transcript")

    add_section_label(document, "听力理解题  /  QUESTIONS")
    question_num_id = new_numbering_instance(document, decimal_abstract)
    for question in sorted(question_set["questions"], key=lambda row: row.get("position", 0)):
        q_type = TYPE_LABELS.get(question.get("type"), question.get("type", ""))
        difficulty = DIFFICULTY_LABELS.get(question.get("difficulty"), question.get("difficulty", ""))
        question_meta = document.add_paragraph(f"{q_type} · {difficulty}", style="Question Meta")
        question_meta.paragraph_format.keep_with_next = True
        prompt = document.add_paragraph(style="Question Prompt")
        apply_numbering(prompt, question_num_id)
        prompt.add_run(question["public"]["prompt"])
        prompt.paragraph_format.keep_with_next = True
        option_num_id = new_numbering_instance(document, option_abstract)
        for option in sorted(question["public"]["options"], key=lambda row: row["id"]):
            paragraph = document.add_paragraph(style="Option")
            apply_numbering(paragraph, option_num_id)
            paragraph.add_run(option["text"])

    add_section_label(document, "答案与中文解析  /  ANSWERS & EXPLANATIONS", answer=True)
    for question in sorted(question_set["questions"], key=lambda row: row.get("position", 0)):
        position = int(question.get("position") or 0)
        answer = question["private"]["answer"].upper()
        lead = document.add_paragraph(style="Answer Lead")
        run = lead.add_run(f"第 {position} 题 · 正确答案 {answer}")
        set_run_font(run, size=11, color=TEAL, bold=True)
        lead.paragraph_format.keep_with_next = True

        p = document.add_paragraph(style="Answer Body")
        label = p.add_run("中文解析：")
        set_run_font(label, size=10.5, color=GOLD, bold=True)
        p.add_run(question["private"]["explanationZh"])

        rationale_num_id = new_numbering_instance(document, option_abstract)
        for option_id in ("a", "b", "c", "d"):
            p = document.add_paragraph(style="Answer Body")
            apply_numbering(p, rationale_num_id)
            p.add_run(question["private"]["optionRationalesZh"][option_id])

        transcript_length = len(transcript)
        for span in question["private"]["evidence"]:
            region, percent = evidence_region(int(span["start"]), transcript_length)
            p = document.add_paragraph(style="Evidence")
            set_paragraph_shading(p, GOLD_LIGHT)
            set_paragraph_left_border(p, GOLD, size=12, space=6)
            location = p.add_run(
                f"原文证据 · {region}（约 {percent}%）· 字符 {span['start']}–{span['end']}\n"
            )
            set_run_font(location, size=8.8, color=GOLD, bold=True)
            quote = p.add_run(f'“{span["quote"]}”')
            set_run_font(quote, size=9.5, color=INK, italic=True)


def ordered_ready_sets(
    library: dict[str, Any], question_bank: dict[str, Any], limit: int | None
) -> tuple[list[dict[str, Any]], dict[str, dict[str, Any]]]:
    items_by_id = {item["id"]: item for item in library["items"]}
    ready = [
        question_set
        for question_set in question_bank["sets"]
        if question_set.get("status") in READY_STATUSES and question_set.get("sourceId") in items_by_id
    ]
    collection_order = {"minute-earth": 0, "bbc-6-minute-english": 1}
    ready.sort(
        key=lambda question_set: (
            collection_order.get(question_set["collection"], 99),
            items_by_id[question_set["sourceId"]].get("year") or 0,
            items_by_id[question_set["sourceId"]].get("sequence") or 0,
            question_set["sourceId"],
        )
    )
    if limit is not None:
        ready = ready[:limit]
    return ready, items_by_id


def build_document(
    library: dict[str, Any], question_bank: dict[str, Any], output: Path, limit: int | None
) -> dict[str, int]:
    ready_sets, items_by_id = ordered_ready_sets(library, question_bank, limit)
    if not ready_sets:
        raise SystemExit("No reviewed question sets were found.")
    document = Document()
    document.core_properties.title = "Aurelis 听力原文、题目与答案解析"
    document.core_properties.subject = "TOEFL-style Academic Listening Practice"
    document.core_properties.author = "Aurelis"
    document.core_properties.keywords = "Aurelis, TOEFL, listening, transcript, questions, answers"
    decimal_abstract, option_abstract = setup_document(document)
    add_cover(document, ready_sets, items_by_id)
    add_overview(document, ready_sets, items_by_id)

    current_collection: str | None = None
    current_segment: str | None = None
    first_content_block = True
    for question_set in ready_sets:
        item = items_by_id[question_set["sourceId"]]
        collection = question_set["collection"]
        starts_new_block = False
        if collection != current_collection:
            collection_heading = document.add_paragraph(COLLECTION_LABELS[collection], style="Heading 1")
            collection_heading.paragraph_format.page_break_before = not first_content_block
            current_collection = collection
            current_segment = None
            starts_new_block = True
        segment = (
            "Minute Earth 已审定题组"
            if collection == "minute-earth"
            else f"{item.get('year') or '未注明年份'} 年"
        )
        if segment != current_segment:
            segment_heading = document.add_paragraph(segment, style="Heading 2")
            segment_heading.paragraph_format.page_break_before = not starts_new_block
            current_segment = segment
            starts_new_block = True
        add_episode(
            document,
            item,
            question_set,
            decimal_abstract,
            option_abstract,
            page_break_before=not starts_new_block,
        )
        first_content_block = False

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "sets": len(ready_sets),
        "questions": len(ready_sets) * 4,
        "minuteEarth": sum(row["collection"] == "minute-earth" for row in ready_sets),
        "bbc": sum(row["collection"] == "bbc-6-minute-english" for row in ready_sets),
    }


def main() -> int:
    args = parse_args()
    counts = build_document(
        load_json(args.library),
        load_json(args.question_bank),
        args.output,
        args.limit,
    )
    print(json.dumps({"output": str(args.output.resolve()), **counts}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
